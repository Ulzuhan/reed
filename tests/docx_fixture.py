"""Build .docx fixtures, including ones python-docx cannot author itself.

Word writes tracked changes as ``w:ins`` and ``w:del`` elements, and python-docx
has no API for them. Splicing them into a generated package is the only way to
test what Reed extracts from a document that is mid-review, which is exactly the
case where extraction can disagree with what a human sees.
"""

from __future__ import annotations

import zipfile
from pathlib import Path

DOCUMENT_PART = "word/document.xml"
_REVISION = ' w:id="{id}" w:author="tester" w:date="2026-01-01T00:00:00Z"'


def write_docx(path: Path, blocks: list[tuple[str, str]]) -> Path:
    """Write a document from ``(style, text)`` pairs, e.g. ``("Heading 1", "…")``."""
    from docx import Document

    document = Document()
    for style, text in blocks:
        document.add_paragraph(text, style=style)
    document.save(str(path))
    return path


def write_docx_with_revisions(path: Path, *, inserted: str, deleted: str) -> Path:
    """Write a document carrying one tracked insertion and one tracked deletion."""
    from docx import Document

    document = Document()
    document.add_paragraph("Unchanged sentence.")
    document.save(str(path))

    with zipfile.ZipFile(path) as archive:
        parts = {name: archive.read(name) for name in archive.namelist()}

    revised = (
        parts[DOCUMENT_PART]
        .decode()
        .replace(
            "</w:body>",
            f"<w:p><w:ins{_REVISION.format(id=1)}>"
            f"<w:r><w:t>{inserted}</w:t></w:r></w:ins></w:p>"
            f"<w:p><w:del{_REVISION.format(id=2)}>"
            f"<w:r><w:delText>{deleted}</w:delText></w:r></w:del></w:p>"
            "</w:body>",
        )
    )

    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as out:
        for name, payload in parts.items():
            out.writestr(name, revised.encode() if name == DOCUMENT_PART else payload)
    return path
